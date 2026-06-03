#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
AI行业商机周报 - Word版生成脚本
覆盖周期：2026.05.16 - 05.22
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import os

doc = Document()

# 全局样式：华文楷体
style = doc.styles['Normal']
font = style.font
font.name = '华文楷体'
style.element.rPr.rFonts.set(qn('w:eastAsia'), '华文楷体')
font.size = Pt(12)

# 段落间距
for level in range(1, 5):
    s = doc.styles[f'Heading {level}']
    s.font.name = '华文楷体'
    s.element.rPr.rFonts.set(qn('w:eastAsia'), '华文楷体')
    pf = s.paragraph_format
    pf.space_before = Pt(12)
    pf.space_after = Pt(6)

# ============ 标题 ============
title = doc.add_heading('AI行业商机周报', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in title.runs:
    run.font.name = '华文楷体'
    run.font.size = Pt(22)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '华文楷体')
    run.font.color.rgb = RGBColor(0, 0, 0)

# 副标题
sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = sub.add_run('覆盖周期：2026年5月16日 — 5月22日')
run.font.size = Pt(12)
run.font.name = '华文楷体'
run.element.rPr.rFonts.set(qn('w:eastAsia'), '华文楷体')

sub2 = doc.add_paragraph()
sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = sub2.add_run('编制：平安银行杭州分行 · 首席行业分析师')
run.font.size = Pt(11)
run.font.name = '华文楷体'
run.element.rPr.rFonts.set(qn('w:eastAsia'), '华文楷体')

sub3 = doc.add_paragraph()
sub3.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = sub3.add_run('资料来源：新华网、36氪、证券时报、财经媒体、浙江省政府公开信息、企查查等')
run.font.size = Pt(10)
run.font.name = '华文楷体'
run.font.color.rgb = RGBColor(128, 128, 128)
run.element.rPr.rFonts.set(qn('w:eastAsia'), '华文楷体')

doc.add_paragraph()  # 空行

# ============ 一、行业动态与发展总结 ============
h1 = doc.add_heading('一、行业动态与发展总结', level=1)

# 1.1 近期政策变动
doc.add_heading('1. 近期政策变动', level=2)

p = doc.add_paragraph()
run = p.add_run('（1）《智能体规范应用与创新发展实施意见》——智能体首次纳入国家战略框架')
run.bold = True

p = doc.add_paragraph('5月8日，国家网信办、国家发展改革委、工业和信息化部联合印发《智能体规范应用与创新发展实施意见》，首次从国家政策层面将智能体定义为具备自主感知、记忆、决策、交互与执行能力的智能系统。文件围绕科学研究、产业发展、提振消费、民生福祉、社会治理等方向，提出19个典型应用场景。这是全球首个以国家政策形式定义智能体产品形态的文件。')

p = doc.add_paragraph()
run = p.add_run('（2）《人工智能终端智能化分级》国家标准——AIoT进入标准化时代')
run.bold = True

p = doc.add_paragraph('同日，工信部、国家标准总局等部门联合发布《人工智能终端智能化分级》（GB/Z 177—2026）系列国家标准，确立L1响应级到L4协同级的四级能力阶梯，首批覆盖手机、电脑、电视、眼镜、汽车座舱、音箱、耳机七大品类。这标志着AI终端能力从模糊宣传转向可测量、可比较、可认证的工程指标体系。')

p = doc.add_paragraph()
run = p.add_run('（3）算力网纳入国家"六张网"基础设施')
run.bold = True

p = doc.add_paragraph('5月中旬，国务院官宣算力网纳入国家"六张网"基础设施（与水、电、通信网并列），2026年总投资超7万亿元。此举标志着国内AI算力基建从单点建设转向全网协同，为大模型、具身智能等产业落地提供关键网络支撑。华为昇腾超节点代表当前国产AI算力高端水平，单套算力达215 PFLOPS，约为英伟达GB200 NVL72超节点的1.2倍。')

p = doc.add_paragraph()
run = p.add_run('（4）《人工智能终端智能化分级》系列国家标准正式发布')
run.bold = True

p = doc.add_paragraph('涉及手机、眼镜、耳机、电脑、电视、汽车座舱、音箱等七大品类的人机交互标准，首次规范AI终端的智能化水平，为终端厂商提供明确的技术路径和认证体系。')

p = doc.add_paragraph()
run = p.add_run('综合影响闭环：')
run.bold = True

p = doc.add_paragraph('以上政策形成"智能体定义+终端分级+算力基建"三位一体的政策矩阵。对AI行业而言，这意味着从技术研发到商业落地的全链条获得了国家级背书，企业可以据此制定3-5年战略规划。银行端应重点关注：一是参与国标起草的企业（华为、荣耀、小米等），具备标准话语权优势；二是19个应用场景落地的AIoT企业，将获得政策补贴和采购倾斜；三是算力网建设带动的服务器、芯片、数据中心等基础设施投资，对应设备融资、项目贷款等金融需求。')

# 1.2 政府规划与产业方向
doc.add_heading('2. 政府规划与产业方向', level=2)

p = doc.add_paragraph('"十五五"规划纲要将机器人与具身智能分别列为新兴产业与未来产业。国务院将算力网纳入国家"六张网"基础设施，2026年总投资超7万亿元。科技部推动新型架构算力系统、人形机器人等重大科技专项。ASIC芯片赛道爆发，2026年ASIC服务器出货增速44.6%，远超GPU服务器的16.1%。全球CSP（谷歌、AWS、微软、Meta）正在统一向自研或定制ASIC倾斜。')

# 1.3 行业发展趋势
doc.add_heading('3. 行业发展趋势', level=2)

p = doc.add_paragraph()
run = p.add_run('（1）大模型"五虎"格局成型，资本集中化加速')
run.bold = True

p = doc.add_paragraph('5月7-8日，中国AI大模型赛道迎来融资决战周：月之暗面完成约20亿美元新融资，投后估值突破200亿美元；阶跃星辰即将完成近25亿美元融资，红筹架构拆除完毕，港股IPO就绪；DeepSeek估值传言飙至500亿美元以上，国家大基金领投。加上港股已上市的智谱（市值约4500亿港元）和MiniMax（市值超3000亿港元），国内基础大模型"五虎"格局已成型。')

p = doc.add_paragraph()
run = p.add_run('（2）商业化拐点出现，AI企业从"烧钱"转向"变现"')
run.bold = True

p = doc.add_paragraph('5月20日的人工智能行业集体业绩说明会上，金山办公、云从科技等企业明确表示AI收入正从试点走向规模化变现。智谱Coding Plan涨价30%仍售罄，MiniMax的M2.5在Agent编程场景供不应求。中金将其总结为"行业定价逻辑从流量消耗转向算力价值变现"。')

p = doc.add_paragraph()
run = p.add_run('（3）ASIC芯片赛道爆发，AI芯片范式转移')
run.bold = True

p = doc.add_paragraph('2026年ASIC服务器出货增速44.6%，远超GPU服务器的16.1%。Cerebras于5月14日在纳斯达克上市，首日大涨89%，市值冲到750亿美元。ASIC的黄金时代本质是AI算力民主化运动，全球AI芯片市场正经历范式转移。')

p = doc.add_paragraph()
run = p.add_run('（4）具身智能从实验室走向产业化')
run.bold = True

p = doc.add_paragraph('具身智能连续两年被写入政府工作报告，杭州具身智能产业集群产值达1068亿元，集聚700余家相关企业。宇树科技发布全球首款量产版载人变形机甲GD01（390万元起），强脑科技发布Revo3灵巧手，标志着具身智能从原型验证迈向规模商业化。')

# 1.4 上下游产业链动态
doc.add_heading('4. 上下游产业链动态', level=2)

p = doc.add_paragraph('上游：AI芯片领域，ASIC赛道火爆，Cerebras上市市值750亿美元。华为昇腾超节点代表国产算力最高水平。A股芯片板块持续活跃，60余家上市公司披露扩产或重大投资计划，覆盖AI算力产业链、半导体等。')

p = doc.add_paragraph('中游：大模型公司集体冲刺商业化，智谱、MiniMax港股表现强劲，月之暗面、阶跃星辰、DeepSeek密集融资。AI企业从"烧钱竞赛"转向"收入验证"阶段。')

p = doc.add_paragraph('下游：AI终端进入标准化时代，L1-L4四级能力阶梯确立。金山办公WPS AI量价齐升，虹软科技布局AI眼镜和车载AI。智能体应用场景扩展至智能制造、智慧金融、智慧能源等领域。')

# ============ 二、浙江地区动态与区域机会 ============
h2 = doc.add_heading('二、浙江地区动态与区域机会', level=1)

doc.add_heading('1. 浙江政策与政府规划', level=2)

p = doc.add_paragraph()
run = p.add_run('（1）国家具身智能中试基地落户杭州')
run.bold = True

p = doc.add_paragraph('5月16日，国家人工智能应用中试基地（具身智能）在杭州滨江区正式揭牌。这是全国唯一面向具身智能领域的国家级应用中试基地，由杭州具身智能中试基地科技有限公司承建（杭州高新区财政局50%控股、杭州市数据集团36%、宇树科技10%、传化智联4%）。基地将构建以算力保障、数据开放、模型服务、场景验证为核心的公共技术服务平台。5月20日，浙江省委常委、杭州市委书记刘非与宇树科技创始人王兴兴等座谈，强调要以超常规举措推进基地建设。')

p = doc.add_paragraph()
run = p.add_run('（2）浙江三大科创高地升级，AI核心产业营收7131亿元')
run.bold = True

p = doc.add_paragraph('5月21日，浙江省科技信息研究院发布三大科创高地2025年度发展报告。2025年，浙江省规上人工智能核心产业营收达7131亿元，数字经济核心产业增加值占GDP比重达13%。新认定高新技术企业4239家，累计3.2万家。宇树科技、群核信息等"杭州六小龙"引起全国关注。面向"十五五"，浙江将聚焦人工智能、生命健康、新材料新能源三大科创高地。')

p = doc.add_paragraph()
run = p.add_run('（3）浙江出台"民营经济新增13条"，AI赋能惠企')
run.bold = True

p = doc.add_paragraph('5月15日，浙江出台"民营经济新增13条"，2026年新增民营经济贷款5000亿元以上，后续余额保持每年5%以上增长。浙江正由省工商联牵头建设"浙里政策惠企精准直达大模型平台"，运用AI技术实现政策与企业精准匹配。')

doc.add_heading('2. 浙江重点产业与区域机会', level=2)

p = doc.add_paragraph('杭州具身智能产业集群产值达1068亿元，集聚700余家相关企业，形成覆盖核心算法、关键零部件、整机制造、场景应用的完整产业链条。杭州出台全国首部聚焦具身智能机器人领域的地方性法规——《杭州市促进具身智能机器人产业发展条例》。')

p = doc.add_paragraph('宁波重点把AI"落子"在制造业领域，发布"工业大模型应用中试工厂"，帮助中小企业解决用不起、不会用AI大模型的问题。')

doc.add_heading('3. 浙江上下游产业链动态', level=2)

p = doc.add_paragraph('上游：之江实验室开发多孔合金专用模型应用于太空计算卫星3D打印；天目山实验室开发航宇陶瓷复材智能体，研发周期缩短30%。')
p = doc.add_paragraph('中游：宇树科技发布载人变形机甲GD01，中科第五纪一个月内完成Pre-A和Pre-A+轮融资（累计数亿元，红杉中国领投），被宇树认证为001号"具身操作大脑"供应商。')
p = doc.add_paragraph('下游：传播大脑与浙江警察学院签署合作协议，聚焦"AI+公安"与公共安全智能治理领域。')

doc.add_heading('4. 银行展业机会', level=2)

p = doc.add_paragraph('1. 具身智能中试基地生态链融资：中试基地共建合伙人体系（宇树科技、传化智联等）及上下游配套企业，存在设备采购、扩产技改、供应链金融需求。')
p = doc.add_paragraph('2. AI大模型商业化变现期授信：浙江本地AI企业进入收入验证阶段，金山办公、虹软科技等上市公司AI业务增长，对应流动资金贷款、票据贴现需求。')
p = doc.add_paragraph('3. AIoT企业扩产融资：A股60余家上市公司扩产，浙江本地企业如传化智联等参与AI生态布局，项目贷款、设备融资租赁需求明确。')
p = doc.add_paragraph('4. 跨境AI业务：虹软科技、云从科技等开拓东南亚、中东市场，对应跨境结算、人民币国际证+福费廷、外币存款需求。')

# ============ 三、重点企业推荐 ============
h3 = doc.add_heading('三、重点企业推荐', level=1)

# 企业1: 宇树科技
doc.add_heading('1. 宇树科技有限公司', level=2)

p = doc.add_paragraph()
run = p.add_run('推荐等级：高')
run.font.name = '华文楷体'
run = p.add_run('\n所属行业：具身智能/机器人\n所在地区：杭州\n产业链位置：整机制造+核心算法\n推荐方向：供应链金融、项目贷款、现金管理、跨境结算')
run.font.name = '华文楷体'

p = doc.add_paragraph()
run = p.add_run('企业关键信息：')
run.bold = True

p = doc.add_paragraph('- 基本情况：成立于2016年，总部杭州，专注于足式机器人和人形机器人研发制造，是国内具身智能赛道头部企业。2025年具身智能产业集群产值1068亿元中占据重要份额。全球首款量产版载人变形机甲GD01（390万元起）于本周发布。')
p = doc.add_paragraph('- 高层与团队背景：创始人兼CEO王兴兴，浙江大学硕士，90后创业者，技术出身，在足式机器人运动控制领域有深厚积累。核心团队来自浙大、清华等高校，在强化学习、运动控制等方向有持续研究产出。')
p = doc.add_paragraph('- 银行关注点：王兴兴参与国家具身智能中试基地创新生态座谈会，在杭州政府层具备较强话语权。企业已赋能近400家机器人客户，生态效应明显。作为中试基地共建合伙人（持股10%），政策资源倾斜确定性强。')

p = doc.add_paragraph()
run = p.add_run('推荐理由：')
run.bold = True

p = doc.add_paragraph('- 具身智能连续两年写入政府工作报告，行业处于爆发前夜，宇树作为国内领先者确定性最高')
p = doc.add_paragraph('- 国家具身智能中试基地落户杭州，宇树为共建合伙人，政策红利直接受益')
p = doc.add_paragraph('- GD01载人机甲发布引发巨大关注（B站播放量近600万），品牌势能强劲')
p = doc.add_paragraph('- 产业链上下游700余家企业集聚杭州，具备批量获客基础')

p = doc.add_paragraph()
run = p.add_run('银行展业机会：')
run.bold = True

p = doc.add_paragraph('- GD01机甲量产带来的设备采购、生产线扩建 → 项目贷款、设备融资、平安租赁')
p = doc.add_paragraph('- 中试基地建设中的配套企业 → 供应链金融、订单融资、付融通')
p = doc.add_paragraph('- 宇树生态链上下游企业 → 国内信用证、银票贴现、商票保贴')

p = doc.add_paragraph()
run = p.add_run('推荐产品组合：')
run.bold = True

p = doc.add_paragraph('建议以项目贷款切入，解决宇树GD01量产线扩建的核心资金需求；配套使用付融通和商票保贴，覆盖中试基地生态链上下游企业的应收账款管理；数字财资用于多主体资金集中管理。具体额度、期限、费率待沟通。')

p = doc.add_paragraph()
run = p.add_run('客户经理切入话术：')
run.bold = True

p = doc.add_paragraph('王总，祝贺宇树GD01机甲发布，关注度非常高。中试基地揭牌后，产业链配套企业肯定会密集落地。我们平安银行在杭州有专门的具身智能服务团队，可以为企业提供从项目建设到供应链的全流程金融方案，方便约个时间当面沟通？')

p = doc.add_paragraph()
run = p.add_run('风险提示：')
run.bold = True

p = doc.add_paragraph('具身智能行业商业化仍在早期，GD01机甲的市场接受度有待验证；核心人才竞争加剧（DeepSeek等公司的人才流失教训）；需结合企业订单、现金流情况综合判断授信额度。')

# 企业2: 虹软科技
doc.add_heading('2. 虹软科技股份有限公司', level=2)

p = doc.add_paragraph()
run = p.add_run('推荐等级：高')
run.font.name = '华文楷体'
run = p.add_run('\n所属行业：AI视觉/智能终端\n所在地区：杭州\n产业链位置：上游算法+中游解决方案\n推荐方向：流动资金贷款、跨境结算、票据贴现、现金管理')
run.font.name = '华文楷体'

p = doc.add_paragraph()
run = p.add_run('企业关键信息：')
run.bold = True

p = doc.add_paragraph('- 基本情况：科创板上市公司（688088），总部杭州，专注于计算机视觉算法和AI视觉解决方案，客户覆盖全球主流智能手机厂商和汽车企业。')
p = doc.add_paragraph('- 高层与团队背景：董事长兼总裁Hui Deng（邓晖），具有深厚的AI视觉领域技术背景和管理经验。')
p = doc.add_paragraph('- 银行关注点：公司战略明确，移动智能终端+车载AI+智能商拍三条业务线并行，正处于AI商业化加速期。')

p = doc.add_paragraph()
run = p.add_run('推荐理由：')
run.bold = True

p = doc.add_paragraph('- AI终端分级标准发布，七大品类覆盖手机、眼镜、汽车座舱等，虹软核心业务直接受益')
p = doc.add_paragraph('- 公司在业绩会上明确表示车载AI业务和智能商拍业务是重点增长方向')
p = doc.add_paragraph('- L1-L4分级体系确立后，AI视觉能力将成为终端产品的核心标尺，虹软作为算法供应商话语权提升')

p = doc.add_paragraph()
run = p.add_run('银行展业机会：')
run.bold = True

p = doc.add_paragraph('- AI眼镜、车载AI等新业务拓展 → 流动资金贷款、科技创新再贷款')
p = doc.add_paragraph('- 海外客户拓展（东南亚、中东等） → 跨境结算、人民币国际证+福费廷')
p = doc.add_paragraph('- 应收账款管理 → 商票e贴、付融通')

p = doc.add_paragraph()
run = p.add_run('推荐产品组合：')
run.bold = True

p = doc.add_paragraph('建议以科技创新和技术更新改造再贷款切入，匹配公司AI算法研发投入的资金需求；配套使用跨境结算和人民币国际证+福费廷，支持海外市场拓展；数字财资用于多币种资金管理。具体参数待沟通。')

p = doc.add_paragraph()
run = p.add_run('客户经理切入话术：')
run.bold = True

p = doc.add_paragraph('邓总，AI终端分级标准落地后，虹软在视觉算法领域的优势会更明显地转化为产品竞争力。我们了解到公司正在大力拓展车载AI和海外市场，平安银行在跨境结算和科创企业融资方面有成熟方案，希望能有机会当面交流。')

p = doc.add_paragraph()
run = p.add_run('风险提示：')
run.bold = True

p = doc.add_paragraph('智能手机市场增长放缓可能影响传统业务；车载AI业务投入期较长，回报存在不确定性；汇率波动影响海外业务利润。仍需结合财报、合同和实地尽调确认。')

# 企业3: 云从科技
doc.add_heading('3. 云从科技集团股份有限公司', level=2)

p = doc.add_paragraph()
run = p.add_run('推荐等级：中')
run.font.name = '华文楷体'
run = p.add_run('\n所属行业：AI大模型/智能体\n所在地区：总部广州，杭州有业务布局\n产业链位置：中游大模型+智能体平台\n推荐方向：订阅模式转型融资、跨境业务、供应链金融')
run.font.name = '华文楷体'

p = doc.add_paragraph()
run = p.add_run('企业关键信息：')
run.bold = True

p = doc.add_paragraph('- 基本情况：科创板上市公司（688327），国内AI四小龙之一，自主研发人机协同操作系统和从容多模态大模型。2025年战略重心全面转向高质量发展。')
p = doc.add_paragraph('- 高层与团队背景：董事长兼总经理周曦，中科院背景，在计算机视觉和AI领域有深厚学术和产业积累。')
p = doc.add_paragraph('- 银行关注点：公司正推动商业模式从一次性交付向订阅与运营服务升级，2026年积极拓展海外市场。')

p = doc.add_paragraph()
run = p.add_run('推荐理由：')
run.bold = True

p = doc.add_paragraph('- 智能体政策落地，云从Agent Lab升级与政策方向高度契合')
p = doc.add_paragraph('- 智慧金融、智能制造等场景已跑通端到端链路，商业化确定性提升')
p = doc.add_paragraph('- 海外市场拓展（东南亚、中东）带来增量空间')

p = doc.add_paragraph()
run = p.add_run('银行展业机会：')
run.bold = True

p = doc.add_paragraph('- 商业模式转型期的营运资金需求 → 流动资金贷款、普惠金融信用贷')
p = doc.add_paragraph('- 海外市场拓展 → 跨境结算、跨境资金管理、外币存款')
p = doc.add_paragraph('- 私有化部署项目 → 项目贷款、设备融资')

p = doc.add_paragraph()
run = p.add_run('推荐产品组合：')
run.bold = True

p = doc.add_paragraph('建议以普惠金融科创贷切入，支持公司智能体平台和Agent Lab升级；配套跨境结算和跨境资金管理方案，匹配东南亚、中东市场拓展需求；资产池用于票据等金融资产全生命周期管理。具体额度、费率待沟通。')

p = doc.add_paragraph()
run = p.add_run('客户经理切入话术：')
run.bold = True

p = doc.add_paragraph('周总，云从在智能体和大模型领域的布局非常扎实，特别是商业模式向订阅升级的方向很清晰。我们平安银行在科创企业服务和跨境业务方面有丰富经验，可以为公司的转型和出海提供支持，方便安排一次拜访吗？')

p = doc.add_paragraph()
run = p.add_run('风险提示：')
run.bold = True

p = doc.add_paragraph('AI商业化转型期收入波动风险；数据安全合规要求在政务、金融行业落地存在不确定性；海外市场拓展面临地缘政治风险。需结合年报、现金流和订单情况综合评估。')

# 企业4: 传化智联
doc.add_heading('4. 传化智联股份有限公司', level=2)

p = doc.add_paragraph()
run = p.add_run('推荐等级：中')
run.font.name = '华文楷体'
run = p.add_run('\n所属行业：智慧物流/产业互联网\n所在地区：杭州\n产业链位置：平台型/场景应用\n推荐方向：项目贷款、供应链金融、现金管理')
run.font.name = '华文楷体'

p = doc.add_paragraph()
run = p.add_run('企业关键信息：')
run.bold = True

p = doc.add_paragraph('- 基本情况：A股上市公司（002010），总部杭州萧山，传化集团旗下核心上市平台，业务涵盖智慧物流、化学制造等。作为股东（持股4%）参与国家具身智能中试基地建设。')
p = doc.add_paragraph('- 高层与团队背景：传化集团创始人徐冠巨，浙商代表人物，在化工和物流领域深耕数十年。')
p = doc.add_paragraph('- 银行关注点：从传统物流向AI+物流转型，中试基地合伙人身份显示其在新制造领域的战略意图。')

p = doc.add_paragraph()
run = p.add_run('推荐理由：')
run.bold = True

p = doc.add_paragraph('- 作为中试基地共建合伙人，深度参与具身智能生态建设')
p = doc.add_paragraph('- 智慧物流场景是具身智能最重要的商业化场景之一')
p = doc.add_paragraph('- 传化在浙江物流网络覆盖广，可作为批量获客的入口')

p = doc.add_paragraph()
run = p.add_run('银行展业机会：')
run.bold = True

p = doc.add_paragraph('- 中试基地配套物流设施建设 → 项目贷款、固定资产贷款')
p = doc.add_paragraph('- 物流平台上下游企业 → 供应链金融、平台数字贷、订货贷')
p = doc.add_paragraph('- 资金结算场景 → 数字财资、慧收款、资产池')

p = doc.add_paragraph()
run = p.add_run('推荐产品组合：')
run.bold = True

p = doc.add_paragraph('建议以项目贷款切入中试基地物流配套建设；配套平台数字贷和订货贷，覆盖物流平台上下游中小商户融资需求；数字财资和慧收款用于平台资金结算管理。具体额度、期限待沟通。')

p = doc.add_paragraph()
run = p.add_run('客户经理切入话术：')
run.bold = True

p = doc.add_paragraph('徐总，传化作为中试基地共建合伙人，在具身智能物流场景的布局很有前瞻性。我们平安银行在供应链金融和物流行业有成熟的产品方案，希望能为传化的数字化转型和生态建设提供资金支持，约个时间详谈？')

p = doc.add_paragraph()
run = p.add_run('风险提示：')
run.bold = True

p = doc.add_paragraph('传统化工业务与AI物流转型的协同性需关注；中试基地投资收益回报周期较长；需结合企业财报和实地调研进一步验证授信可行性。')

# ============ 四、客户经理行动建议 ============
h4 = doc.add_heading('四、客户经理行动建议', level=1)

p = doc.add_paragraph()
run = p.add_run('1. 优先拜访宇树科技及中试基地生态链企业')
run.bold = True
p = doc.add_paragraph('国家具身智能中试基地已于5月16日揭牌，宇树科技为共建合伙人。建议本周内对接宇树科技融资需求（GD01量产线扩建），同时摸排中试基地周边700余家机器人企业的金融需求，以供应链金融和票据产品批量切入。')

p = doc.add_paragraph()
run = p.add_run('2. 以AI终端分级标准为契机，摸排虹软科技等AIoT企业')
run.bold = True
p = doc.add_paragraph('L1-L4分级体系确立，AI视觉算法企业将直接受益。建议对接虹软科技，以科创贷+跨境结算方案切入，同时关注宁波"工业大模型应用中试工厂"相关企业的设备融资需求。')

p = doc.add_paragraph()
run = p.add_run('3. 关注大模型商业化拐点，布局AI企业订阅转型融资')
run.bold = True
p = doc.add_paragraph('智谱Coding Plan售罄、MiniMax Agent场景限购，说明AI付费拐点已现。建议摸排浙江本地AI企业的订阅转型融资需求，以普惠金融科创贷和流动资金贷款切入。')

p = doc.add_paragraph()
run = p.add_run('4. 准备具身智能行业专属金融方案')
run.bold = True
p = doc.add_paragraph('具身智能是浙江重点培育产业集群，建议内部联动产品经理，形成针对具身智能企业的标准化金融方案包：设备融资（平安租赁）+ 供应链金融（付融通/商票保贴）+ 现金管理（数字财资）+ 跨境服务。')

p = doc.add_paragraph()
run = p.add_run('5. 提前布局算力网投资机遇')
run.bold = True
p = doc.add_paragraph('国务院将算力网纳入"六张网"基础设施，2026年投资超7万亿元。建议提前摸排浙江本地数据中心、算力基础设施建设企业的融资需求，以项目贷款和银团融资切入。')

# 保存
output_path = '/Users/leidongqiao/.openclaw/workspace/workspace-AIResearcher/reports/ai-weekly/行业商机周报_AI_20260522.docx'
doc.save(output_path)
print(f'Saved to {output_path}')
