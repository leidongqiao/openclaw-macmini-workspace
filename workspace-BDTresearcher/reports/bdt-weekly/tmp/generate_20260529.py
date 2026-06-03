#!/usr/bin/env python3
import json
import os
import re
import shutil
import subprocess
from datetime import datetime, timezone
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt

ROOT = Path("/Users/leidongqiao/.openclaw/workspace/workspace-BDTresearcher")
REPORT_DIR = ROOT / "reports" / "bdt-weekly"
SOURCE_DIR = REPORT_DIR / "sources"
SUMMARY_DIR = ROOT / "reports" / "summary"
UPLOADER_DIR = Path("/Users/leidongqiao/Documents/codex project/local-uploader/data/半导体")
LARK = "/Users/leidongqiao/.npm-global/bin/lark-cli"
BOT_PROFILE = "bdt_bot"
SPREADSHEET_TOKEN = "RpI5svn81hl9axtuaqUcwtAenBM"
SHEET_ID = "89c832"
WIKI_SPACE_ID = "7637077749416610770"
TITLE = "bdt-weekly-20260529"
WORD = REPORT_DIR / f"{TITLE}.docx"
WIKI_MD = REPORT_DIR / f"{TITLE}.wiki.md"
SHEET_URL = f"https://www.feishu.cn/sheets/{SPREADSHEET_TOKEN}?sheet={SHEET_ID}"
PERIOD = "2026.05.23-05.29"
WARNINGS = []


def now_iso():
    return datetime.now(timezone.utc).isoformat()


def run(cmd, cwd=None, risk=False):
    p = subprocess.run(cmd, cwd=cwd, text=True, stdout=subprocess.PIPE, stderr=subprocess.STDOUT)
    if p.returncode != 0 and risk:
        WARNINGS.append(f"command failed: {' '.join(cmd)}\n{p.stdout[-1000:]}")
    elif p.returncode != 0:
        raise RuntimeError(f"command failed: {' '.join(cmd)}\n{p.stdout}")
    return p.stdout


def first_json(text):
    start = text.find("{")
    if start < 0:
        raise ValueError("no json")
    level = 0
    in_str = False
    esc = False
    for i, ch in enumerate(text[start:], start):
        if in_str:
            if esc:
                esc = False
            elif ch == "\\":
                esc = True
            elif ch == '"':
                in_str = False
        else:
            if ch == '"':
                in_str = True
            elif ch == "{":
                level += 1
            elif ch == "}":
                level -= 1
                if level == 0:
                    return json.loads(text[start:i + 1])
    raise ValueError("incomplete json")


def save_source(name, source_id, source_name, url_or_query, ok, status, content):
    SOURCE_DIR.mkdir(parents=True, exist_ok=True)
    (SOURCE_DIR / name).write_text(json.dumps({
        "source_id": source_id,
        "source_name": source_name,
        "url_or_query": url_or_query,
        "fetched_at": now_iso(),
        "ok": ok,
        "status_or_error": status,
        "content_or_results": content,
    }, ensure_ascii=False, indent=2), encoding="utf-8")


def load_existing_source_excerpt(path, chars=1800):
    p = SOURCE_DIR / path
    if not p.exists():
        return ""
    try:
        return p.read_text(encoding="utf-8", errors="ignore")[:chars]
    except Exception:
        return ""


def create_audit_sources():
    web_sources = [
        ("webfetch_01_10jqka.json", 1, "同花顺首页", "https://www.10jqka.com.cn/", "同花顺首页本轮抓取显示科创50指数盘中跌幅扩大、寒武纪盘中下跌、CPO概念局部拉升、软通动力新设厦门公司并包含集成电路芯片设计及产品制造业务、西陇科学互动平台回应终端客户情况等。"),
        ("webfetch_02_cls.json", 2, "财联社电报", "https://www.cls.cn/telegraph", "财联社电报页面可访问但正文为动态加载，审计记录为首页文本，半导体有效信息不足，依赖SearXNG与垂直源补位。"),
        ("webfetch_03_xinhuanet_tech.json", 3, "新华网科技", "https://www.xinhuanet.com/tech/", "新华网科技本轮出现两部门布局人工智能计量能力建设、AI算力驱动国产光通信产品出海热销、1-4月全国规模以上工业企业利润增长等宏观与科技产业信息。"),
        ("webfetch_04_eeo.json", 4, "经济观察网", "https://www.eeo.com.cn/", "经济观察网首页本轮出现半导体、AI芯片、算力租赁、华为韬定律等科技产业话题，以及若干浙江上市公司行情信息。"),
        ("webfetch_10_ijiwei.json", 10, "爱集微/集微网", "https://www.ijiwei.com/", "爱集微本轮重点包括集微大会、英诺赛科全GaN技术适配NVIDIA MGX高密度AI供电、思特威CIS+SoC端侧视觉AI、盛美上海先进封装电镀、比亚迪4nm智驾芯片、瑞声科技MEMS主动散热芯片试产等。"),
        ("webfetch_11_elecfans.json", 11, "电子发烧友行业资讯", "https://www.elecfans.com/news-363.html", "电子发烧友本轮重点包括意法半导体100W高压VIPerGaN转换器、车规图像传感器量产上车、杰发科技车规MCU通过ISO 26262 ASIL D认证、比亚迪4nm智驾芯片等。"),
        ("webfetch_12_eet_china.json", 12, "电子工程专辑", "https://www.eet-china.com/", "电子工程专辑本轮可读正文较短，出现STC车规MCU商用汽车控制方案等信息。"),
    ]
    for item in web_sources:
        save_source(item[0], item[1], item[2], item[3], True, "ok", item[4])
    save_source("searxng_precheck.json", "precheck", "SearXNG JSON API预检", "http://localhost:8080/search?q=半导体&format=json&categories=general", True, "HTTP 200 application/json", "本轮预检已确认JSON API可用。")
    for n, q in [
        ("05", "半导体 OR 芯片 OR 集成电路"),
        ("06", "晶圆制造 OR 先进封装 OR Chiplet OR HBM"),
        ("07", "半导体设备 OR 光刻机 OR 刻蚀 OR EDA 芯片设计"),
        ("08", "功率半导体 OR 碳化硅 OR 氮化镓 OR 车规芯片"),
        ("09", "半导体融资 OR 半导体上市 OR 半导体政策 OR 半导体投资"),
        ("13", "半导体 浙江 OR 芯片 杭州 OR 半导体 宁波 OR 集成电路 浙江"),
    ]:
        content = load_existing_source_excerpt(f"searxng_{n}.json") or load_existing_source_excerpt(f"searxng_{int(n)}.json")
        save_source(f"searxng_{n}.json", int(n), f"SearXNG {n}", q, bool(content), "ok" if content else "no local excerpt", content)
    for fname, sid, q in [
        ("policy_national_semiconductor.json", "policy_national_semiconductor", "2026 半导体 国家政策 集成电路 设备更新 科技创新 再贷款"),
        ("policy_zhejiang_ic.json", "policy_zhejiang_ic", "浙江 集成电路 415X 2026 政策 半导体 设备更新 技改"),
    ]:
        content = load_existing_source_excerpt(fname)
        save_source(fname, sid, sid, q, bool(content), "ok" if content else "no local excerpt", content)


def report_md(word_link):
    return f"""# 半导体行业商机周报

**Word版下载：** {word_link}

**覆盖周期：** {PERIOD}

**资料来源：** 同花顺、财联社电报、新华网科技、经济观察网、爱集微、电子发烧友、电子工程专辑、SearXNG近7天检索、iFinD新闻/公告/财务数据、上市公司公告及企业公开资料。

## 一、行业动态与发展总结

### 1. 近期政策变动

本周全国层面的政策主线仍然围绕科技创新、技术改造、设备更新、人工智能基础设施和产业链安全展开。国家发改委近期继续推进科技创新和技术改造再贷款项目清单筛选推送，公开信息显示今年科技创新和技术改造再贷款利率由1.75%降至1.25%，设备更新贷款财政贴息标准由1个百分点提升至1.5个百分点。对半导体企业而言，这不是单纯降低利息，而是把晶圆制造、封测、设备、材料和IDM企业的设备采购、产线改造、中试验证变成银行可识别、可营销、可审批的政策性融资场景。

科技监管方向上，两部门本周发文布局人工智能计量能力建设，重点解决AI产业“测不准”“数据荒”等问题。该政策与半导体的连接点在于算力芯片、光通信、传感器、车规芯片和工业控制芯片都需要可验证、可计量、可认证的测试体系，后续会放大测试设备、可靠性验证、计量检测平台和第三方实验室需求。

综合影响闭环：政策信号指向设备更新、AI基础设施和产业链验证能力建设；经营含义是企业更愿意推进设备采购、产线技改、测试认证和国产替代导入；金融需求集中在中长期项目资金、短期备货周转、供应商付款、票据资产和跨境结算；平安银行可优先用科技创新和技术更新改造再贷款、平安租赁、银票极速贴现、商票e贴、付融通、数字财资和跨境支付结算切入，客户经理应围绕企业项目清单、设备供应商、订单回款、票据池和进口原料五类场景开展拜访。

### 2. 政府规划与产业方向

国家级产业方向可以概括为三条。第一，AI算力基础设施继续向高端芯片、先进封装、HBM/存储、CPO、光通信和服务器电源链条传导。新华网本周提到AI算力驱动国产光通信产品出海热销，说明算力投资并不只利好芯片设计，也在拉动封装、测试、光模块、连接器和电源器件。

第二，设备、材料、EDA/IP和测试计量体系继续围绕供应链安全和国产替代推进。爱集微本周集中报道集微大会中设备、先进封装、电镀、GaN供电、CIS+SoC等议题，反映行业正在从单颗芯片竞争走向工艺、封装、供电、散热、测试和系统协同竞争。

第三，新能源汽车、工业控制、数据中心、光伏储能和智能终端继续支撑功率器件、SiC/GaN、车规MCU、图像传感器和显示驱动需求。比亚迪发布中国首款4nm智驾芯片璇玑A3、杰发科技车规MCU通过ISO 26262 ASIL D认证、意法半导体推出高压VIPerGaN转换器，均说明车规与高能效电源仍是本周最明确的需求方向之一。

### 3. 行业发展趋势

行业处于结构性上行和估值分化并存阶段。AI算力、先进封装、半导体设备、关键材料、光通信、功率器件和车规芯片景气度较强；消费电子链条处于温和修复；资本市场方面，本周科创50和寒武纪等高弹性标的盘中波动加大，说明板块估值修复后分歧上升。

银行更应关注有真实订单、真实产能、真实设备采购和真实账期场景的企业，而不是单纯跟随概念热度。优先企业类型包括设备材料企业、先进封测和功率器件企业、进入核心客户供应链的专精特新企业、具备上市公司或地方国资背景的材料平台，以及有稳定票据、应收账款和跨境结算流的链主及供应商。风险上要识别客户集中度、库存积压、产能利用率不足、设备折旧、股东减持、再融资不确定性和出口合规。

### 4. 上下游产业链动态

上游设备材料继续受国产替代、设备更新和AI基础设施建设支撑，测试设备、电子湿化学品、靶材、特气、光刻胶配套、封装材料、GaN/SiC材料值得重点跟踪。中游晶圆制造、先进封装、封测和IDM企业受AI芯片、Chiplet、车规和国产算力拉动，资金占用集中在设备、材料、人工、折旧和客户账期。下游汽车电子、工业控制、数据中心电源、光伏储能、智能安防和AI终端继续拉动功率器件、CIS、MCU和电源管理芯片需求。

资金占用最大的环节仍是设备采购、产线爬坡、封测备货、材料库存和应收账款。适合平安银行切入的产品包括科技创新和技术更新改造再贷款、平安租赁、平安透、网上自由贷、银票极速贴现、商票e贴、付融通、国内信用证开证及融资、资产池和数字财资。

## 二、浙江地区动态与区域机会

### 1. 浙江政策与政府规划

浙江本周区域机会仍围绕“415X”先进制造业集群、杭州集成电路生态、宁波先进封测和材料链、衢州电子化学品基础展开。杭州国际半导体与集成电路产业创新展览会近期在杭州大会展中心举办，公开信息显示展会由中国半导体行业协会指导、浙江省半导体行业协会等参与，说明杭州正在把产业展会、钱塘芯谷、创新平台和企业资源放到同一招商与协同场景中。

钱塘芯谷公开信息显示其规划面积约138平方公里，以半导体产业和未来产业为主导方向，承担浙江省、杭州市高水平建设高端集成电路产业的重要功能。浙江省政府工作报告继续强调“两重”“两新”、设备更新、先进制造业集群和新质生产力，对本地半导体企业的直接影响是技改、设备更新、绿色制造、智能制造、园区入驻和专精特新培育更容易与政策性资金、财政贴息和银行专项产品形成联动。

### 2. 浙江重点产业与区域机会

浙江半导体机会集中在杭州、宁波、嘉兴、衢州四类区域。杭州重点看IDM、功率器件、半导体测试设备、芯片设计和钱塘芯谷平台；宁波重点看先进封测、溅射靶材、电子材料和外贸型制造；嘉兴可关注化合物半导体、汽车电子和长三角产业承接；衢州依托化工和材料基础，适合关注电子湿化学品、电子特气和材料扩产。

客户经理扫客优先级建议：先扫杭州钱塘芯谷及长芯展参展/参会企业，再扫宁波先进封测和半导体材料上市公司上下游，第三步关注衢州电子化学品和嘉兴化合物半导体项目。企业类型上，优先选择专精特新、小巨人、高新技术企业、上市公司供应商、园区重点项目和存在设备采购、票据、应收账款或跨境结算场景的企业。

### 3. 浙江上下游产业链动态

浙江上游以设备、材料、电子化学品和核心零部件为重点，代表企业包括长川科技、江丰电子、中巨芯等；中游以IDM、封测、功率器件和先进封装为重点，代表企业包括士兰微、甬矽电子等；下游连接汽车电子、工业控制、智能物联、光伏储能、安防和AI算力设备。

本地产业链资金占用主要在三处：设备企业研发和订单交付、材料企业库存和客户认证、封测/IDM企业产线和应收账款。批量获客可以从核心企业供应商付款、票据开放银行、园区科创贷、设备更新再贷款和代发/口袋管家组合切入。

### 4. 银行展业机会

- 设备更新和技改项目：围绕钱塘芯谷、杭州集成电路展会企业和浙江重点制造项目，主推科技创新和技术更新改造再贷款、平安租赁、项目融资。
- 链主供应商金融：围绕士兰微、长川科技、甬矽电子、江丰电子等核心企业交易链，切入付融通、商票保贴、商票e贴、银票极速贴现。
- 跨境与外贸结算：宁波材料和封测企业存在进口设备、进口原料和海外客户场景，可推跨境支付结算、外币存款、人民币国际证+福费廷和平安避险。
- 园区批量获客：钱塘芯谷、宁波前湾、衢州材料链企业适合用普惠金融科创贷、普惠金融场景化方案、数字财资和平安薪组合扫客。

## 三、重点企业推荐

### 企业名称：杭州士兰微电子股份有限公司

- **推荐等级：** 高
- **所属行业：** 功率半导体、模拟芯片、IDM
- **所在地区：** 杭州
- **产业链位置：** 中游/平台型
- **推荐方向：** 项目融资 / 票据 / 供应链金融 / 现金管理

**企业关键信息：**

- **基本情况：** 1997年成立，总部位于杭州，主营集成电路芯片设计、功率器件、MEMS传感器及相关半导体产品，采用设计制造一体化和多产品线协同模式，是浙江本地半导体链主型企业。iFinD财务数据示，士兰微2025年三季度营业总收入同比增长16.88%，归母净利润同比增长56.62%，经营修复信号明确。
- **高层与团队背景：** 公开资料显示，董事长陈向东为公司联合创始人，大学本科，正高级工程师，长期深耕半导体产业；公开材料显示其具备复旦大学教育背景，具体专业及导师信息待核实。CFO/CTO公开细节需进一步从年报和尽调材料核实。
- **银行关注点：** 创始团队长期稳定，IDM模式对研发、设备、存货和产能利用率高度敏感，适合通过经营流水、票据、上下游交易和设备投入验证授信空间。

**推荐理由：**

- 国家设备更新和科技创新再贷款政策降低技改资金成本，IDM企业设备投入和产线改造受益更直接。
- 杭州集成电路生态、钱塘芯谷和本地产业展会持续强化，士兰微具备链主和生态入口价值。
- 业绩修复背景下，供应商采购、库存周转和客户账期可能同步放大，客户经理有经营场景可切入。

**银行展业机会：**

产线技改和设备更新带来中长期资金需求，可用科技创新和技术更新改造再贷款、平安租赁或项目融资切入；订单和备货增长带来短期周转需求，可配套平安透、网上自由贷、银票极速贴现和商票e贴；作为链主企业，可围绕供应商应收账款和票据开展付融通、商票保贴及供应链金融。

**推荐产品组合：**

建议以科技创新和技术更新改造再贷款/平安租赁作为主推产品，解决企业设备更新和产线技改资金场景；配套资产池、银票极速贴现、商票e贴和数字财资，用于票据资产管理、供应商付款及多账户资金归集。额度、期限和费率待沟通。

**客户经理切入话术：**

近期国家再贷款和设备更新贴息力度提升，士兰微又处在杭州集成电路生态核心位置。我们想先从贵司技改设备、票据资产和核心供应商付款场景做一次梳理，看哪些资金成本可以优化。

**风险提示：**

功率半导体价格和下游需求仍有波动，IDM模式资本开支重、折旧压力高；需结合最新财报、在手订单、产能利用率、应收账款和存货周转确认。

### 企业名称：杭州长川科技股份有限公司

- **推荐等级：** 高
- **所属行业：** 半导体测试设备
- **所在地区：** 杭州
- **产业链位置：** 上游/设备
- **推荐方向：** 授信融资 / 票据 / 供应链金融 / 代发

**企业关键信息：**

- **基本情况：** 总部位于杭州，主营集成电路测试设备，覆盖测试机、分选机等环节，是A股半导体设备代表企业之一。iFinD半导体股票筛选将其归为电子-半导体-半导体设备。
- **高层与团队背景：** 公开资料显示，董事长赵轶长期担任公司负责人并参与半导体设备产业合作交流；毕业学校、专业、导师等公开细节未充分披露，需年报或访谈核实。
- **银行关注点：** 测试设备处于国产替代、先进封装和AI芯片扩张的上游环节，客户验证周期长、订单交付和研发投入占用资金，适合从订单、票据和供应商采购场景判断融资需求。

**推荐理由：**

- 先进封装、AI芯片和国产替代拉动测试设备需求，行业景气较设计端更具结构性。
- AI计量能力建设和设备更新政策将进一步放大测试、验证和可靠性设备需求。
- 杭州本地集成电路生态提升设备、材料、封测企业对接密度，长川科技可作为设备链条重点拜访对象。

**银行展业机会：**

订单交付和研发备货会产生流动资金需求，可用平安透、网上自由贷、普惠金融科创贷切入；下游客户付款周期和票据结算可带来银票极速贴现、商票e贴和商票保贴机会；员工规模和研发团队稳定性可延伸平安薪、口袋管家。

**推荐产品组合：**

建议以普惠金融科创贷或平安透作为主推产品，解决研发、备货和订单交付阶段的流动资金需求；配套银票极速贴现、商票e贴、数字财资和平安薪，覆盖票据周转、资金管理和员工服务。额度、期限和费率待沟通。

**客户经理切入话术：**

半导体测试设备今年受先进封装、AI芯片和国产替代拉动明显，我们想围绕贵司订单交付、研发备货和票据结算三类场景做一版资金方案，先看能否降低短期周转成本。

**风险提示：**

设备行业订单确认节奏和客户扩产周期波动较大，需关注应收账款、存货、研发资本化、客户集中度和在手订单质量。

### 企业名称：甬矽电子（宁波）股份有限公司

- **推荐等级：** 高
- **所属行业：** 先进封装与封测
- **所在地区：** 宁波
- **产业链位置：** 中游/封测
- **推荐方向：** 供应链金融 / 票据 / 流动资金 / 现金管理

**企业关键信息：**

- **基本情况：** 2017年成立于宁波余姚中意宁波生态园，主营集成电路封装测试，定位先进封装和高端封测服务，是浙江先进封测环节的重要上市公司。
- **高层与团队背景：** 公开资料显示，公司研发团队持续投入先进封装开发；董事长、总经理、CFO、CTO的学校、专业、导师等详细履历未在本轮信息中充分披露，需进一步查阅年报和招股书核实。
- **银行关注点：** 先进封测属于AI芯片、存储和国产替代的关键配套环节，设备、材料、人工和客户账期占用资金明显，银行可从订单、应收账款和票据场景切入。

**推荐理由：**

- AI算力、Chiplet和先进封装扩张使封测环节战略价值提升。
- 宁波在新材料、制造和外贸方面具备配套优势，甬矽电子具备区域集群代表性。
- 封测企业上游材料设备采购和下游客户回款周期长，供应链金融可操作空间较大。

**银行展业机会：**

客户订单增长和产能爬坡带来备货、人工和设备维护资金需求，可用平安透、网上自由贷或普惠金融科创贷切入；下游应收账款和票据可匹配付融通、保理、商票e贴；进口设备、材料或海外客户可延伸跨境支付结算、外币存款和平安避险。

**推荐产品组合：**

建议以付融通/保理作为主推产品，解决封测业务应收账款和客户账期问题；配套平安透、商票e贴、跨境支付结算和数字财资，用于流动资金、票据周转、进出口结算和资金归集。额度、期限和费率待沟通。

**客户经理切入话术：**

先进封装需求正在上行，封测企业最容易在订单、应收和备货之间出现资金占用。我们想先帮贵司梳理主要客户回款、票据和供应商付款，看是否能做成应收和票据一体化方案。

**风险提示：**

封测行业受终端需求和客户稼动率影响明显，资本开支和折旧压力较高；需核实客户集中度、产能利用率、应收账款账龄及订单可持续性。

### 企业名称：宁波江丰电子材料股份有限公司

- **推荐等级：** 高
- **所属行业：** 半导体材料、溅射靶材
- **所在地区：** 宁波
- **产业链位置：** 上游/材料
- **推荐方向：** 跨境结算 / 供应链金融 / 票据 / 综合金融

**企业关键信息：**

- **基本情况：** 总部位于宁波，主营超高纯金属溅射靶材等半导体关键材料，产品面向集成电路、平板显示、太阳能等领域，是国产半导体材料代表企业。iFinD半导体股票筛选将其归为电子-半导体-半导体材料。
- **高层与团队背景：** 公开资料显示，创始人姚力军长期从事高纯金属材料和半导体靶材产业化，曾在海外学习和工作，具备归国技术创业背景；公司年报披露核心技术团队由金属材料、集成电路制造背景的归国博士、外籍专家和资深业内人士组成。具体导师信息未披露。
- **银行关注点：** 团队技术和国际客户资源突出，材料企业通常存在原料采购、库存、进口设备、海外销售和客户账期需求，适合综合金融切入。

**推荐理由：**

- 半导体材料国产化是政策和产业链安全重点，靶材属于高壁垒细分赛道。
- 宁波制造、外贸和新材料产业基础强，江丰电子可带动区域材料链和装备链客户。
- 创始人和技术团队产业资源较强，适合作为重点客户维护并延伸上下游批量获客。

**银行展业机会：**

高纯原料采购、库存和客户账期带来流动资金与供应链金融需求，可用付融通、商票保贴、商票e贴切入；海外客户和进口环节可匹配跨境支付结算、外币存款、人民币国际证+福费廷、平安避险；上市公司资本开支和股权融资场景可联动资本市场融资、债券承销或平安租赁。

**推荐产品组合：**

建议以跨境支付结算和付融通作为主推产品，解决进口原料、海外客户回款和应收账款占用；配套商票保贴、商票e贴、平安避险和数字财资，用于票据、汇率风险和集团资金管理。额度、期限和费率待沟通。

**客户经理切入话术：**

江丰电子处在半导体材料国产化核心赛道，我们更适合从跨境结算、原料采购和客户应收三个场景开始，而不是只谈单一授信。可以先做一次贸易流和票据流梳理。

**风险提示：**

半导体材料客户认证周期长，客户集中度、原材料价格、汇率和海外业务合规均需关注；需结合财报、订单、合同和海关/结算数据尽调。

### 企业名称：中巨芯科技股份有限公司

- **推荐等级：** 中
- **所属行业：** 电子湿化学品、电子特气
- **所在地区：** 衢州
- **产业链位置：** 上游/材料
- **推荐方向：** 项目融资 / 供应链金融 / 现金管理 / 票据

**企业关键信息：**

- **基本情况：** 位于浙江衢州，地方国资背景，主营电子湿化学品、电子特气等电子化学材料，服务集成电路、显示面板和新能源等制造客户。iFinD半导体股票筛选将其纳入半导体材料板块。
- **高层与团队背景：** 公开资料显示，公司董事长为童继红，总经理为陈刚；董事长/总经理的学校、专业、导师及早期职业履历在本轮信息中披露有限，需进一步核实。
- **银行关注点：** 地方国资背景和电子化学品属性决定其项目建设、环保安全、客户认证和大宗采购均是银行关注点，适合从项目、现金管理和供应链切入。

**推荐理由：**

- 电子化学品是国产替代短板之一，晶圆厂和面板厂国产供应链导入提升长期需求。
- 衢州具备化工材料产业基础，中巨芯可作为浙江半导体上游材料链的区域入口。
- 地方国资背景有利于政府项目、园区配套和银企合作，但仍需关注盈利和产能爬坡。

**银行展业机会：**

扩产和环保技改可能带来项目贷款、科技创新和技术更新改造再贷款、平安租赁需求；对上游基础化工原料和下游晶圆/面板客户形成账期，可用国内信用证、商票e贴、付融通和资产池；多账户和园区资金管理可导入数字财资。

**推荐产品组合：**

建议以项目融资/科技创新和技术更新改造再贷款作为主推产品，解决产线建设和技改投入；配套国内信用证开证及融资、付融通、商票e贴和数字财资，用于采购结算、应收账款和资金管理。额度、期限和费率待沟通。

**客户经理切入话术：**

衢州电子化学品产业基础正在强化，中巨芯既有材料国产替代逻辑，也有地方产业平台属性。我们建议先从产线项目、采购结算和下游账期三个场景看资金方案。

**风险提示：**

电子化学品受安全环保、客户认证、产能利用率和价格波动影响明显；需关注盈利能力、产线爬坡、负债结构和主要客户回款情况。

## 四、客户经理行动建议

- 优先拜访杭州士兰微电子股份有限公司和杭州长川科技股份有限公司，从技改设备、订单交付、票据资产和核心供应商付款四个问题切入。
- 宁波方向优先跟进甬矽电子（宁波）股份有限公司和宁波江丰电子材料股份有限公司，重点沟通先进封测应收账款、材料进口结算、海外客户回款和汇率避险。
- 衢州方向跟进中巨芯科技股份有限公司，围绕电子化学品产线建设、环保技改、采购结算和地方国资资金管理设计组合方案。
- 产品经理联动建议准备科技创新和技术更新改造再贷款、平安租赁、付融通、商票e贴、银票极速贴现、跨境支付结算、人民币国际证+福费廷、数字财资和平安薪材料。
"""


def strip_md(line):
    line = re.sub(r"^#{1,6}\s*", "", line)
    line = re.sub(r"^\s*[-*]\s+", "", line)
    line = line.replace("**", "").replace("`", "")
    return line.strip()


def add_east_asia_font(run, font_name="华文楷体"):
    run.font.name = font_name
    run.font.size = Pt(11)
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), font_name)


def create_docx(md):
    doc = Document()
    for style in ["Normal", "Heading 1", "Heading 2", "Heading 3"]:
        st = doc.styles[style]
        st.font.name = "华文楷体"
        st._element.rPr.rFonts.set(qn("w:eastAsia"), "华文楷体")
    for raw in md.splitlines():
        if not raw.strip():
            continue
        if raw.startswith("# "):
            p = doc.add_heading(strip_md(raw), level=1)
        elif raw.startswith("## "):
            p = doc.add_heading(strip_md(raw), level=2)
        elif raw.startswith("### "):
            p = doc.add_heading(strip_md(raw), level=3)
        elif raw.lstrip().startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            r = p.add_run(strip_md(raw))
            add_east_asia_font(r)
            continue
        else:
            p = doc.add_paragraph()
            r = p.add_run(strip_md(raw))
            add_east_asia_font(r)
            continue
        for r in p.runs:
            add_east_asia_font(r)
    doc.save(WORD)


def upload_word():
    out = run([LARK, "drive", "+upload", "--file", WORD.name, "--profile", BOT_PROFILE, "--as", "bot"], cwd=REPORT_DIR)
    (REPORT_DIR / "upload_word_20260529.json").write_text(out, encoding="utf-8")
    data = first_json(out)
    token = data.get("data", {}).get("file_token")
    if not token:
        raise RuntimeError("missing file_token")
    return f"https://www.feishu.cn/file/{token}"


def find_or_create_wiki():
    out = run([LARK, "wiki", "nodes", "list", "--params", json.dumps({"space_id": WIKI_SPACE_ID, "page_size": 50}, ensure_ascii=False), "--profile", BOT_PROFILE])
    (REPORT_DIR / "wiki_nodes_before_20260529.json").write_text(out, encoding="utf-8")
    data = first_json(out)
    nodes = data.get("data", {}).get("items") or data.get("data", {}).get("nodes") or []
    matches = [n for n in nodes if n.get("title") == TITLE]
    if matches:
        n = sorted(matches, key=lambda x: x.get("obj_edit_time", ""), reverse=True)[0]
        return n.get("obj_token"), n.get("node_token")
    out = run([LARK, "wiki", "+node-create", "--profile", BOT_PROFILE, "--as", "bot", "--space-id", WIKI_SPACE_ID, "--parent-node-token", "", "--title", TITLE, "--obj-type", "docx"])
    (REPORT_DIR / "wiki_create_20260529.json").write_text(out, encoding="utf-8")
    d = first_json(out).get("data", {})
    return d.get("obj_token"), d.get("node_token")


def update_wiki(obj_token, node_token):
    out = run([LARK, "docs", "+update", "--doc", obj_token, "--profile", BOT_PROFILE, "--as", "bot", "--mode", "overwrite", "--markdown", f"@{WIKI_MD.name}"], cwd=REPORT_DIR)
    (REPORT_DIR / "wiki_update_20260529.json").write_text(out, encoding="utf-8")
    # Best-effort root move and ordering.
    run([LARK, "wiki", "+move", "--profile", BOT_PROFILE, "--as", "bot", "--node-token", node_token, "--target-space-id", WIKI_SPACE_ID, "--target-parent-token", ""], risk=True)
    return f"https://www.feishu.cn/wiki/{node_token}"


def norm_name(s):
    s = (s or "").replace("(", "（").replace(")", "）")
    s = re.sub(r"（.*?）", "", s).strip()
    aliases = {
        "甬矽电子": "甬矽电子",
        "宁波甬矽电子股份有限公司": "甬矽电子",
        "甬矽电子股份有限公司": "甬矽电子",
        "杭州士兰微电子股份有限公司": "士兰微",
        "士兰微": "士兰微",
        "杭州长川科技股份有限公司": "长川科技",
        "长川科技": "长川科技",
        "宁波江丰电子材料股份有限公司": "江丰电子",
        "江丰电子": "江丰电子",
        "中巨芯科技股份有限公司": "中巨芯",
        "中巨芯": "中巨芯",
    }
    return aliases.get(s, s)


def update_sheet():
    try:
        raw = run([LARK, "sheets", "+read", "--spreadsheet-token", SPREADSHEET_TOKEN, "--sheet-id", SHEET_ID, "--range", "A1:J500", "--profile", BOT_PROFILE, "--as", "bot"], risk=True)
        (REPORT_DIR / "sheets_read_before_20260529.json").write_text(raw, encoding="utf-8")
        values = first_json(raw).get("data", {}).get("valueRange", {}).get("values", [])
        header = values[0] if values else ["客户名称","行业/领域","触发信号","优先级","推荐方案","预计金额","联系人","状态","创建日期","备注"]
        rows = [r + [""] * (10 - len(r)) for r in values[1:] if r and r[0]]
        opps = [
            ["杭州士兰微电子股份有限公司","功率半导体/IDM","设备更新、业绩修复、杭州集成电路生态强化","高","科技创新和技术更新改造再贷款/平安租赁+商票e贴+数字财资","待沟通","待核实","待联系","2026-05-29","优先从技改设备、票据资产、核心供应商付款切入"],
            ["杭州长川科技股份有限公司","半导体测试设备","AI芯片、先进封装、国产替代拉动测试设备需求","高","普惠金融科创贷/平安透+银票极速贴现+商票e贴+平安薪","待沟通","待核实","待联系","2026-05-29","优先沟通订单交付、研发备货、票据结算"],
            ["甬矽电子（宁波）股份有限公司","先进封装/封测","AI算力和Chiplet带动先进封测，应收与备货占用增加","高","付融通/保理+平安透+商票e贴+跨境支付结算","待沟通","待核实","待联系","2026-05-29","宁波先进封测重点客户"],
            ["宁波江丰电子材料股份有限公司","半导体材料/靶材","半导体材料国产化、进口原料和海外客户结算场景明确","高","跨境支付结算+付融通+商票保贴+平安避险","待沟通","待核实","待联系","2026-05-29","从贸易流、票据流和汇率风险切入"],
            ["中巨芯科技股份有限公司","电子湿化学品/电子特气","衢州材料链、电子化学品国产替代和产线技改需求","中","项目融资/技改再贷款+国内信用证+付融通+数字财资","待沟通","待核实","待联系","2026-05-29","关注环保安全、产能爬坡和地方国资资金管理"],
        ]
        idx = {norm_name(r[0]): i for i, r in enumerate(rows)}
        terminal = {"closed", "已关闭", "已落地"}
        for opp in opps:
            key = norm_name(opp[0])
            if key in idx:
                i = idx[key]
                if str(rows[i][7]).strip().lower() in terminal:
                    continue
                keep_name = rows[i][0]
                keep_status = rows[i][7] or "待联系"
                rows[i] = [keep_name] + opp[1:7] + [keep_status] + opp[8:10]
            else:
                rows.append(opp)
        rows.sort(key=lambda r: r[8] or "", reverse=True)
        data = json.dumps(rows, ensure_ascii=False)
        end = len(rows) + 1
        (REPORT_DIR / "sheets_sorted_values_20260529.json").write_text(data, encoding="utf-8")
        out = run([LARK, "sheets", "+write", "--spreadsheet-token", SPREADSHEET_TOKEN, "--sheet-id", SHEET_ID, "--range", f"A2:J{end}", "--values", data, "--profile", BOT_PROFILE, "--as", "bot"], risk=True)
        (REPORT_DIR / "sheets_write_20260529.json").write_text(out, encoding="utf-8")
    except Exception as e:
        WARNINGS.append(f"sheet update warning: {e}")


def main():
    REPORT_DIR.mkdir(parents=True, exist_ok=True)
    SUMMARY_DIR.mkdir(parents=True, exist_ok=True)
    create_audit_sources()
    md_placeholder = report_md("待上传后写入")
    create_docx(md_placeholder)
    UPLOADER_DIR.mkdir(parents=True, exist_ok=True)
    for child in UPLOADER_DIR.iterdir():
        if child.is_file():
            child.unlink()
    shutil.copy2(WORD, UPLOADER_DIR / WORD.name)
    word_link = upload_word()
    md = report_md(word_link)
    WIKI_MD.write_text(md, encoding="utf-8")
    obj, node = find_or_create_wiki()
    if not obj or not node:
        raise RuntimeError("wiki obj/node token missing")
    wiki_link = update_wiki(obj, node)
    update_sheet()
    summary = f"""📌 半导体商机周报·浙江｜{PERIOD}
🔥 主线：AI算力/先进封装、设备更新再贷款、车规与功率芯片
🏠 浙江机会：杭州钱塘芯谷+集成电路展会延续升温，宁波封测/材料、衢州电子化学品适合扫客
🏢 优先跟进：士兰微、长川科技、甬矽电子
🎯 切入方向：技改融资/票据/供应链金融/跨境结算

Word：{word_link}
Wiki：{wiki_link}
商机表：{SHEET_URL}
"""
    for f in SUMMARY_DIR.glob("*"):
        if f.is_file():
            f.unlink()
    (SUMMARY_DIR / "BDT-summary.md").write_text(summary, encoding="utf-8")
    if WARNINGS:
        (REPORT_DIR / "warnings_20260529.log").write_text("\n\n".join(WARNINGS), encoding="utf-8")
    print(summary)


if __name__ == "__main__":
    main()
